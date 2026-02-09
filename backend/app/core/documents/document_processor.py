"""Document Processor for RAG Knowledge Base."""
import os
import hashlib
import asyncio
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process documents and extract text for RAG."""
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md', '.html', '.csv']
    
    async def process_document(
        self,
        file_path: str,
        document_id: str,
        document_type: str = "general"
    ) -> Dict[str, Any]:
        """
        Process a document and return chunks for indexing.
        
        Args:
            file_path: Path to the document
            document_id: ID of the document in database
            document_type: Type of document (faq, manual, guide, etc.)
        
        Returns:
            Dict with 'chunks' and 'metadata'
        """
        try:
            logger.info(f"[DOC_PROCESS] Starting to process document: {document_id}")
            logger.info(f"[DOC_PROCESS] File path: {file_path}")
            
            # Extract text based on file type
            logger.info(f"[DOC_PROCESS] Extracting text from file...")
            text = await self.extract_text(file_path)
            
            logger.info(f"[DOC_PROCESS] Extracted {len(text)} characters")
            
            if not text.strip():
                logger.warning(f"[DOC_PROCESS] No text extracted from document {document_id}")
                return {
                    'success': False,
                    'error': 'No text extracted from document',
                    'chunks': []
                }
            
            logger.info(f"[DOC_PROCESS] First 100 chars: {text[:100]}...")
            
            # Generate content hash
            content_hash = hashlib.md5(text.encode()).hexdigest()
            logger.info(f"[DOC_PROCESS] Content hash: {content_hash}")
            
            # Chunk the text
            logger.info(f"[DOC_PROCESS] Chunking text...")
            chunks = self.chunk_text(text)
            
            logger.info(f"[DOC_PROCESS] Created {len(chunks)} chunks")
            
            # Add metadata to each chunk
            chunk_data = []
            for i, chunk in enumerate(chunks):
                chunk_data.append({
                    'content': chunk,
                    'document_id': document_id,
                    'document_type': document_type,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'content_hash': content_hash
                })
                if i < 3:  # Log first 3 chunks
                    logger.info(f"[DOC_PROCESS] Chunk {i}: {chunk[:100]}...")
            
            logger.info(f"[DOC_PROCESS] Completed processing document {document_id}: {len(chunks)} chunks")
            
            return {
                'success': True,
                'content_hash': content_hash,
                'chunks': chunk_data,
                'chunk_count': len(chunks)
            }
            
        except Exception as e:
            logger.error(f"[DOC_PROCESS] Error processing document {document_id}: {e}")
            import traceback
            logger.error(f"[DOC_PROCESS] Traceback: {traceback.format_exc()}")
            return {
                'success': False,
                'error': str(e),
                'chunks': []
            }
    
    async def extract_text(self, file_path: str) -> str:
        """Extract text from various document formats."""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return await self.extract_from_pdf(file_path)
        elif ext == '.docx':
            return await self.extract_from_docx(file_path)
        elif ext in ['.txt', '.md']:
            return await self.extract_from_text(file_path)
        elif ext == '.html':
            return await self.extract_from_html(file_path)
        elif ext == '.csv':
            return await self.extract_from_csv(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    async def extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        logger.info(f"[PDF] Starting PDF extraction for: {file_path}")
        
        try:
            import PyPDF2
            logger.info(f"[PDF] PyPDF2 imported successfully")
            
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                logger.info(f"[PDF] PDF has {len(reader.pages)} pages")
                
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    text += page_text
                    if i == 0:  # Log first page
                        logger.info(f"[PDF] Page 0: {page_text[:200]}...")
            
            logger.info(f"[PDF] Total extracted text length: {len(text)}")
            return text
            
        except ImportError:
            logger.warning("[PDF] PyPDF2 not installed, attempting alternative extraction...")
            # Try to use pdfminer or other available libraries
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract
                text = pdfminer_extract(file_path)
                logger.info(f"[PDF] Used pdfminer, extracted {len(text)} chars")
                return text
            except ImportError:
                logger.warning("[PDF] No PDF library available (PyPDF2, pdfminer)")
                logger.warning("[PDF] Installing PyPDF2 is recommended: pip install PyPDF2")
                # Return empty text - document won't be indexed
                return ""
        except Exception as e:
            logger.error(f"[PDF] Error extracting from PDF: {e}")
            import traceback
            logger.error(f"[PDF] Traceback: {traceback.format_exc()}")
            return ""
    
    async def extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text.append(" | ".join(row_text))
            
            return "\n".join(text)
            
        except ImportError:
            logger.warning("python-docx not installed")
            return ""
    
    async def extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text or markdown."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    async def extract_from_html(self, file_path: str) -> str:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator='\n')
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except ImportError:
            logger.warning("beautifulsoup4 not installed")
            return await self.extract_from_text(file_path)
    
    async def extract_from_csv(self, file_path: str) -> str:
        """Extract text from CSV."""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            text = df.to_string()
            return text
            
        except ImportError:
            # Fallback to basic CSV parsing
            lines = []
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    lines.append(line.strip())
            return '\n'.join(lines)
    
    async def extract_text_simple(self, file_path: str) -> str:
        """Simple text extraction as fallback."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            return ""
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if not text:
            return []
        
        # Check if text is mostly binary (PDF headers, etc.)
        # If text contains many non-printable characters, it's likely binary
        non_printable_ratio = sum(1 for c in text if not (c.isprintable() or c in '\n\t\r')) / len(text)
        if non_printable_ratio > 0.3:
            logger.warning(f"[CHUNK] Text appears to be binary ({non_printable_ratio*100:.1f}% non-printable), returning empty chunks")
            return []
        
        # Limit text size to prevent memory issues
        max_text_length = 100000  # 100KB max
        if len(text) > max_text_length:
            logger.warning(f"[CHUNK] Text too long ({len(text)} chars), truncating to {max_text_length} chars")
            text = text[:max_text_length]
        
        chunks = []
        start = 0
        text_length = len(text)
        max_chunks = 100  # Limit total chunks to prevent memory issues
        
        while start < text_length and len(chunks) < max_chunks:
            end = min(start + self.chunk_size, text_length)
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence endings
                for sep in ['. ', '! ', '? ', '\n']:
                    last_sep = text.rfind(sep, start, end)
                    if last_sep > start + self.chunk_size // 2:
                        end = last_sep + len(sep)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start, accounting for overlap
            start = end - self.chunk_overlap
        
        logger.info(f"[CHUNK] Created {len(chunks)} chunks from {text_length} chars")
        return chunks
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        import re
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?;:\'-]', '', text)
        
        return text.strip()
    
    async def process_text_content(
        self,
        content: str,
        document_id: str,
        document_type: str,
        title: str = ""
    ) -> Dict[str, Any]:
        """Process raw text content (not from file)."""
        # Clean the text
        content = self.clean_text(content)
        
        if not content.strip():
            return {
                'success': False,
                'error': 'No text content to process',
                'chunks': []
            }
        
        # Generate content hash
        content_hash = hashlib.md5(content.encode()).hexdigest()
        
        # Chunk the text
        chunks = self.chunk_text(content)
        
        # Add metadata to each chunk
        chunk_data = []
        for i, chunk in enumerate(chunks):
            chunk_data.append({
                'content': chunk,
                'document_id': document_id,
                'document_type': document_type,
                'title': title,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'content_hash': content_hash
            })
        
        return {
            'success': True,
            'content_hash': content_hash,
            'chunks': chunk_data,
            'chunk_count': len(chunks)
        }


# Singleton instance
document_processor = DocumentProcessor()
